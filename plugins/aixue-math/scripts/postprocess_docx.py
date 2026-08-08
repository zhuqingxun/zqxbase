# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx>=1.1"]
# ///
"""postprocess_docx.py — pandoc 输出的 docx 二次加工

解决 pandoc 3.x 产出教案 docx 的三个典型问题：

1. 表格缺边框：为每个 cell 显式加 tcBorders（全边框 single 4sz）
2. 中文字体提示缺失：所有 run 的 w:rFonts 补 w:eastAsia="Microsoft YaHei"
3. 粗体不清：bold run 强制字体名 Microsoft YaHei（确保 Word 伪粗渲染清晰）

用法：
    uv run --script postprocess_docx.py <input.docx> [<output.docx>]

若省略 output：原地改写 input.docx。
"""
from __future__ import annotations

import argparse
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


EAST_ASIAN_FONT = "Microsoft YaHei"
ASCII_FONT_FOR_BOLD = "Microsoft YaHei"  # bold run 也用雅黑
MONO_FONT = "Consolas"


def add_cell_borders(tc) -> None:
    """给单元格添加全边框。"""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        tcBorders.append(e)
    tcPr.append(tcBorders)


def add_table_borders(tbl) -> None:
    """给整个表格加全边框（含内部 insideH/insideV）。"""
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        tblBorders.append(e)
    tblPr.append(tblBorders)


def ensure_east_asian_font(run_element, east_asia_font: str = EAST_ASIAN_FONT) -> None:
    """给 run 的 w:rPr 下补齐 w:rFonts 的 eastAsia 属性。"""
    rPr = run_element.find(qn("w:rPr"))
    if rPr is None:
        return
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    if not rFonts.get(qn("w:eastAsia")):
        rFonts.set(qn("w:eastAsia"), east_asia_font)
    if not rFonts.get(qn("w:ascii")):
        rFonts.set(qn("w:ascii"), east_asia_font)
    if not rFonts.get(qn("w:hAnsi")):
        rFonts.set(qn("w:hAnsi"), east_asia_font)


def force_bold_font(run_element, bold_font: str = ASCII_FONT_FOR_BOLD) -> None:
    """若 run 是 bold，强制其字体为指定字体（确保伪粗渲染清晰）。"""
    rPr = run_element.find(qn("w:rPr"))
    if rPr is None:
        return
    b = rPr.find(qn("w:b"))
    if b is None:
        return
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), bold_font)
    rFonts.set(qn("w:hAnsi"), bold_font)
    rFonts.set(qn("w:eastAsia"), bold_font)
    rFonts.set(qn("w:cs"), bold_font)


def is_code_run(run_element) -> bool:
    """判断 run 是否属于代码/等宽样式（Verbatim Char 或 Source Code）。"""
    rPr = run_element.find(qn("w:rPr"))
    if rPr is None:
        return False
    rStyle = rPr.find(qn("w:rStyle"))
    if rStyle is None:
        return False
    style_val = rStyle.get(qn("w:val")) or ""
    return style_val in ("VerbatimChar", "SourceCode")


def process_document(doc: Document) -> dict:
    stats = {"tables": 0, "cells": 0, "runs": 0, "bold_runs": 0, "code_runs": 0}

    # 表格：整体 + 单元格
    for tbl in doc.element.body.iter(qn("w:tbl")):
        stats["tables"] += 1
        add_table_borders(tbl)
        for tc in tbl.iter(qn("w:tc")):
            stats["cells"] += 1
            add_cell_borders(tc)

    # 所有 run：字体提示 + 粗体强化
    for r in doc.element.body.iter(qn("w:r")):
        stats["runs"] += 1
        if is_code_run(r):
            stats["code_runs"] += 1
            continue  # 代码 run 保留 Consolas
        ensure_east_asian_font(r)
        rPr = r.find(qn("w:rPr"))
        if rPr is not None and rPr.find(qn("w:b")) is not None:
            stats["bold_runs"] += 1
            force_bold_font(r)

    return stats


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("input_docx")
    ap.add_argument("output_docx", nargs="?", default=None)
    args = ap.parse_args()

    inp = Path(args.input_docx).resolve()
    if not inp.is_file():
        print(f"ERROR: 输入不存在: {inp}", file=sys.stderr)
        sys.exit(1)

    out = Path(args.output_docx).resolve() if args.output_docx else inp
    if out != inp:
        shutil.copyfile(inp, out)

    doc = Document(out)
    stats = process_document(doc)
    doc.save(out)

    print(f"postprocessed {out}")
    print(f"  tables={stats['tables']} cells={stats['cells']} runs={stats['runs']} "
          f"bold_runs={stats['bold_runs']} code_runs_kept={stats['code_runs']}")


if __name__ == "__main__":
    main()
